import re
from pathlib import Path

import pandas as pd

from docx import Document as DocxDocument

from pypdf import PdfReader

from models import Document


class FileReader:

    SUPPORTED_EXTENSIONS = {
        ".txt",
        ".csv",
        ".pdf",
        ".docx",
    }

    # ========================================================
    # DIRECTORY
    # ========================================================

    def read_directory(
            self,
            directory: Path,
    ) -> list[Document]:

        if not directory.exists():

            raise FileNotFoundError(
                "Data folder does not exist:\n"
                f"{directory.absolute()}"
            )

        documents = []

        for path in sorted(
                directory.iterdir()
        ):

            if not path.is_file():
                continue

            extension = (
                path.suffix.lower()
            )

            if (
                    extension
                    not in self.SUPPORTED_EXTENSIONS
            ):
                continue

            try:

                content = self.read_file(
                    path
                )

                content = self.clean_text(
                    content
                )

                if not content:
                    continue

                documents.append(
                    Document(
                        filename=path.name,
                        content=content,
                    )
                )

                print(
                    f"Loaded: {path.name}"
                )

            except Exception as error:

                print(
                    f"Could not read "
                    f"{path.name}: {error}"
                )

        return documents

    # ========================================================
    # FILE
    # ========================================================

    def read_file(
            self,
            path: Path,
    ) -> str:

        readers = {

            ".txt":
                self._read_txt,

            ".csv":
                self._read_csv,

            ".pdf":
                self._read_pdf,

            ".docx":
                self._read_docx,
        }

        reader = readers.get(
            path.suffix.lower()
        )

        if reader is None:

            raise ValueError(
                f"Unsupported file type: "
                f"{path.suffix}"
            )

        return reader(path)

    # ========================================================
    # TXT
    # ========================================================

    @staticmethod
    def _read_txt(
            path: Path,
    ) -> str:

        return path.read_text(
            encoding="utf-8",
            errors="ignore",
        )

    # ========================================================
    # CSV
    # ========================================================

    @staticmethod
    def _read_csv(
            path: Path,
    ) -> str:

        dataframe = pd.read_csv(
            path
        )

        return dataframe.to_string(
            index=False
        )

    # ========================================================
    # PDF
    # ========================================================

    @staticmethod
    def _read_pdf(
            path: Path,
    ) -> str:

        reader = PdfReader(path)

        pages = []

        for page_number, page in enumerate(
                reader.pages,
                start=1,
        ):

            page_text = (
                page.extract_text()
            )

            if page_text:

                pages.append(
                    f"[Page {page_number}]\n"
                    f"{page_text}"
                )

        return "\n\n".join(
            pages
        )

    # ========================================================
    # DOCX
    # ========================================================

    @staticmethod
    def _read_docx(
            path: Path,
    ) -> str:

        document = DocxDocument(
            path
        )

        parts = []

        # ----------------------------------------------------
        # Paragraphs
        # ----------------------------------------------------

        for paragraph in (
                document.paragraphs
        ):

            text = (
                paragraph.text.strip()
            )

            if text:
                parts.append(text)

        # ----------------------------------------------------
        # Tables
        # ----------------------------------------------------

        for table_number, table in enumerate(
                document.tables,
                start=1,
        ):

            parts.append(
                f"[TABLE {table_number}]"
            )

            for row in table.rows:

                cells = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                parts.append(
                    " | ".join(cells)
                )

        return "\n".join(parts)

    # ========================================================
    # CLEANING
    # ========================================================

    @staticmethod
    def clean_text(
            text: str,
    ) -> str:

        text = re.sub(
            r"[ \t]+",
            " ",
            text,
        )

        text = re.sub(
            r"\n{3,}",
            "\n\n",
            text,
        )

        return text.strip()