from pathlib import Path
import hashlib
import json
import re
import time

import httpx
import numpy as np
import pandas as pd
from pypdf import PdfReader
from docx import Document


# ============================================================
# CONFIGURATION
# ============================================================

DATA_FOLDER = Path("../data")

# Folder used to store the local RAG index
INDEX_FOLDER = Path("../index")

# Ollama API
OLLAMA_URL = "http://localhost:11434"


# ============================================================
# LOCAL LLM
# ============================================================

LLM_MODEL = "qwen3:1.7b"

# Keep Qwen loaded for 10 minutes.
# This makes the second/third question much faster.
LLM_KEEP_ALIVE = "10m"


# ============================================================
# EMBEDDING MODEL
# ============================================================

EMBEDDING_MODEL = "nomic-embed-text:latest"

# Keep embedding model loaded for 5 minutes.
# This avoids reloading it for every question.
EMBEDDING_KEEP_ALIVE = "5m"


# ============================================================
# RAG SETTINGS
# ============================================================

# Number of chunks returned by semantic search
TOP_K = 2

# Chunk size in characters
CHUNK_SIZE = 1000

# Overlap between chunks
CHUNK_OVERLAP = 150


# ============================================================
# PERFORMANCE SETTINGS
# ============================================================

# Small batch because the computer has only 12 GB RAM
EMBED_BATCH_SIZE = 2

# Context size for Qwen
NUM_CTX = 1024

# Maximum generated answer length
# Lower value = faster/shorter responses.
NUM_PREDICT = 80

# Your CPU has 2 physical cores
OLLAMA_NUM_THREADS = 4


# ============================================================
# CACHE FILES
# ============================================================

CHUNKS_FILE = INDEX_FOLDER / "chunks.json"

EMBEDDINGS_FILE = INDEX_FOLDER / "embeddings.npy"

CACHE_INFO_FILE = INDEX_FOLDER / "cache_info.json"


# ============================================================
# HTTP CLIENT
# ============================================================

# Reuse HTTP connections instead of creating a new connection
# for every Ollama request.
HTTP_CLIENT = httpx.Client(
    timeout=httpx.Timeout(
        connect=10.0,
        read=300.0,
        write=300.0,
        pool=10.0
    )
)


# ============================================================
# OLLAMA
# ============================================================

def check_ollama():
    """
    Check whether Ollama is running.
    """

    try:

        response = HTTP_CLIENT.get(
            f"{OLLAMA_URL}/api/tags"
        )

        response.raise_for_status()

        return True

    except Exception:

        return False


def get_ollama_models():
    """
    Return installed Ollama models.
    """

    try:

        response = HTTP_CLIENT.get(
            f"{OLLAMA_URL}/api/tags"
        )

        response.raise_for_status()

        data = response.json()

        return [
            model["name"]
            for model in data.get("models", [])
        ]

    except Exception:

        return []


# ============================================================
# FILE READERS
# ============================================================

def read_txt(file_path):
    """
    Read a TXT file.
    """

    with open(
            file_path,
            "r",
            encoding="utf-8",
            errors="ignore"
    ) as file:

        return file.read()


def read_csv(file_path):
    """
    Read a CSV file.
    """

    dataframe = pd.read_csv(
        file_path
    )

    return dataframe.to_string(
        index=False
    )


def read_pdf(file_path):
    """
    Extract text from PDF.
    """

    reader = PdfReader(
        file_path
    )

    text = []

    for page_number, page in enumerate(
            reader.pages,
            start=1
    ):

        page_text = page.extract_text()

        if page_text:

            text.append(
                f"[Page {page_number}]\n"
                f"{page_text}"
            )

    return "\n\n".join(text)


def read_docx(file_path):
    """
    Read DOCX paragraphs and tables.
    """

    document = Document(
        file_path
    )

    text = []

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        paragraph_text = paragraph.text.strip()

        if paragraph_text:

            text.append(
                paragraph_text
            )

    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    for table_number, table in enumerate(
            document.tables,
            start=1
    ):

        text.append(
            f"\n[TABLE {table_number}]"
        )

        for row in table.rows:

            cells = []

            for cell in row.cells:

                cells.append(
                    cell.text.strip()
                )

            text.append(
                " | ".join(cells)
            )

    return "\n".join(text)


# ============================================================
# CONVERT FILE TO TEXT
# ============================================================

def convert_file_to_text(file_path):

    extension = file_path.suffix.lower()

    try:

        if extension == ".txt":

            return read_txt(
                file_path
            )

        elif extension == ".csv":

            return read_csv(
                file_path
            )

        elif extension == ".pdf":

            return read_pdf(
                file_path
            )

        elif extension == ".docx":

            return read_docx(
                file_path
            )

        return None

    except Exception as error:

        print(
            f"Could not read "
            f"{file_path.name}: {error}"
        )

        return None


# ============================================================
# LOAD DOCUMENTS
# ============================================================

def load_documents():

    documents = []

    if not DATA_FOLDER.exists():

        print(
            "ERROR: Data folder does not exist:"
        )

        print(
            f"    {DATA_FOLDER.absolute()}"
        )

        return documents

    for file_path in sorted(
            DATA_FOLDER.iterdir()
    ):

        if not file_path.is_file():

            continue

        text = convert_file_to_text(
            file_path
        )

        if text is None:

            continue

        text = text.strip()

        if not text:

            continue

        documents.append({

            "filename": file_path.name,

            "content": text

        })

    return documents


# ============================================================
# TEXT CLEANING
# ============================================================

def clean_text(text):

    # Replace multiple spaces
    text = re.sub(
        r"[ \t]+",
        " ",
        text
    )

    # Replace excessive newlines
    text = re.sub(
        r"\n{3,}",
        "\n\n",
        text
    )

    return text.strip()


# ============================================================
# CHUNK DOCUMENTS
# ============================================================

def create_chunks(documents):

    chunks = []

    for document in documents:

        filename = document["filename"]

        content = clean_text(
            document["content"]
        )

        start = 0

        content_length = len(content)

        while start < content_length:

            end = min(
                start + CHUNK_SIZE,
                content_length
            )

            chunk_text = content[
                start:end
            ].strip()

            if chunk_text:

                chunks.append({

                    "filename": filename,

                    "content": chunk_text

                })

            if end >= content_length:

                break

            start = end - CHUNK_OVERLAP

    return chunks


# ============================================================
# CACHE / INDEX
# ============================================================

def calculate_documents_hash(documents):
    """
    Create a hash based on document contents and
    important RAG configuration.

    If documents or configuration change,
    embeddings will be recreated.
    """

    hasher = hashlib.sha256()

    # Include configuration in the hash

    hasher.update(
        str(CHUNK_SIZE).encode()
    )

    hasher.update(
        str(CHUNK_OVERLAP).encode()
    )

    hasher.update(
        EMBEDDING_MODEL.encode()
    )

    # Include every document

    for document in documents:

        hasher.update(
            document["filename"].encode(
                "utf-8",
                errors="ignore"
            )
        )

        hasher.update(
            document["content"].encode(
                "utf-8",
                errors="ignore"
            )
        )

    return hasher.hexdigest()


def save_index(
        chunks,
        embeddings,
        documents_hash
):
    """
    Save chunks and embeddings to disk.
    """

    INDEX_FOLDER.mkdir(
        parents=True,
        exist_ok=True
    )

    # --------------------------------------------------------
    # Save chunks
    # --------------------------------------------------------

    with open(
            CHUNKS_FILE,
            "w",
            encoding="utf-8"
    ) as file:

        json.dump(
            chunks,
            file,
            ensure_ascii=False,
            indent=2
        )

    # --------------------------------------------------------
    # Save embeddings
    # --------------------------------------------------------

    np.save(
        EMBEDDINGS_FILE,
        embeddings
    )

    # --------------------------------------------------------
    # Save cache information
    # --------------------------------------------------------

    cache_info = {

        "documents_hash":
            documents_hash,

        "embedding_model":
            EMBEDDING_MODEL,

        "chunk_size":
            CHUNK_SIZE,

        "chunk_overlap":
            CHUNK_OVERLAP,

        "number_of_chunks":
            len(chunks),

        "embedding_dimensions":
            int(embeddings.shape[1])
            if embeddings.ndim == 2
            else 0
    }

    with open(
            CACHE_INFO_FILE,
            "w",
            encoding="utf-8"
    ) as file:

        json.dump(
            cache_info,
            file,
            indent=2
        )


def load_index(documents_hash):
    """
    Load cached chunks and embeddings if they
    match the current documents/configuration.
    """

    if not (
            CHUNKS_FILE.exists()
            and EMBEDDINGS_FILE.exists()
            and CACHE_INFO_FILE.exists()
    ):

        return None, None

    try:

        # ----------------------------------------------------
        # Load cache information
        # ----------------------------------------------------

        with open(
                CACHE_INFO_FILE,
                "r",
                encoding="utf-8"
        ) as file:

            cache_info = json.load(
                file
            )

        # ----------------------------------------------------
        # Check document hash
        # ----------------------------------------------------

        if (
                cache_info.get("documents_hash")
                != documents_hash
        ):

            return None, None

        # ----------------------------------------------------
        # Check embedding model
        # ----------------------------------------------------

        if (
                cache_info.get("embedding_model")
                != EMBEDDING_MODEL
        ):

            return None, None

        # ----------------------------------------------------
        # Check chunk configuration
        # ----------------------------------------------------

        if (
                cache_info.get("chunk_size")
                != CHUNK_SIZE
        ):

            return None, None

        if (
                cache_info.get("chunk_overlap")
                != CHUNK_OVERLAP
        ):

            return None, None

        # ----------------------------------------------------
        # Load chunks
        # ----------------------------------------------------

        with open(
                CHUNKS_FILE,
                "r",
                encoding="utf-8"
        ) as file:

            chunks = json.load(
                file
            )

        # ----------------------------------------------------
        # Load embeddings
        # ----------------------------------------------------

        embeddings = np.load(
            EMBEDDINGS_FILE
        )

        # ----------------------------------------------------
        # Validate embedding array
        # ----------------------------------------------------

        if embeddings.ndim != 2:

            return None, None

        if len(chunks) != len(embeddings):

            return None, None

        if len(chunks) == 0:

            return None, None

        return (
            chunks,
            embeddings
        )

    except Exception as error:

        print(
            f"Could not load cache: {error}"
        )

        return None, None


# ============================================================
# OLLAMA EMBEDDINGS
# ============================================================

def create_embeddings(texts):

    response = httpx.post(
        f"{OLLAMA_URL}/api/embed",
        json={
            "model": EMBEDDING_MODEL,
            "input": texts,
            "keep_alive": 0
        },
        timeout=300
    )

    response.raise_for_status()

    data = response.json()

    embeddings = data.get("embeddings")

    if not embeddings:
        raise RuntimeError(
            "Ollama did not return embeddings."
        )

    return np.array(
        embeddings,
        dtype=np.float32
    )

# ============================================================
# CREATE EMBEDDINGS FOR ALL CHUNKS
# ============================================================

def embed_chunks(chunks):

    print()

    print(
        f"Creating embeddings for "
        f"{len(chunks)} chunks..."
    )

    texts = [
        chunk["content"]
        for chunk in chunks
    ]

    embeddings = []

    total = len(texts)

    start_time = time.perf_counter()

    for start in range(
            0,
            total,
            EMBED_BATCH_SIZE
    ):

        batch = texts[
            start:
            start + EMBED_BATCH_SIZE
        ]

        end_number = min(
            start + EMBED_BATCH_SIZE,
            total
        )

        print(
            f"\rEmbedding "
            f"{end_number}/{total}...",
            end="",
            flush=True
        )

        batch_embeddings = create_embeddings(
            batch
        )

        embeddings.extend(
            batch_embeddings
        )

    elapsed = (
            time.perf_counter()
            - start_time
    )

    print()

    print(
        f"Embedding completed in "
        f"{elapsed:.1f} seconds."
    )

    return np.array(
        embeddings,
        dtype=np.float32
    )

# ----------------------------------------------------
# Greeting SEARCH
# ----------------------------------------------------
def is_greeting(question):
    greetings = {
        "hi",
        "hello",
        "hey",
        "thanks",
        "thank you",
        "bye",
        "exit",
        "good morning",
        "good afternoon",
        "good evening"
    }

    normalized = question.lower().strip()

    return normalized in greetings

# ============================================================
# COSINE SIMILARITY
# ============================================================

def semantic_search(
        question,
        chunks,
        embeddings,
        top_k=TOP_K
):
    """
    Perform semantic search using vectorized
    NumPy cosine similarity.
    """

    start_time = time.perf_counter()

    # --------------------------------------------------------
    # Create embedding for the question
    # --------------------------------------------------------

    question_embedding = create_embeddings(
        [question]
    )[0]

    # --------------------------------------------------------
    # Normalize question vector
    # --------------------------------------------------------

    question_norm = np.linalg.norm(
        question_embedding
    )

    if question_norm == 0:

        return []

    normalized_question = (
            question_embedding
            / question_norm
    )

    # --------------------------------------------------------
    # Normalize document embeddings
    # --------------------------------------------------------

    embedding_norms = np.linalg.norm(
        embeddings,
        axis=1,
        keepdims=True
    )

    embedding_norms[
        embedding_norms == 0
        ] = 1.0

    normalized_embeddings = (
            embeddings
            / embedding_norms
    )

    # --------------------------------------------------------
    # Calculate all similarities at once
    # --------------------------------------------------------

    scores = np.dot(
        normalized_embeddings,
        normalized_question
    )

    # --------------------------------------------------------
    # Get top results
    # --------------------------------------------------------

    top_k = min(
        top_k,
        len(scores)
    )

    top_indices = np.argsort(
        scores
    )[::-1][:top_k]

    results = []

    for index in top_indices:

        result = {

            "filename":
                chunks[index]["filename"],

            "content":
                chunks[index]["content"],

            "score":
                float(scores[index])
        }

        results.append(
            result
        )

    elapsed = (
            time.perf_counter()
            - start_time
    )

    print(
        f"Semantic search completed in "
        f"{elapsed:.2f} seconds."
    )

    return results


# ============================================================
# ASK LOCAL LLM
# ============================================================

def ask_ai(
        question,
        search_results
):

    context_parts = []

    for result in search_results:

        context_parts.append(

            "SOURCE: "
            + result["filename"]
            + "\n"
            + result["content"]

        )

    context = "\n\n".join(
        context_parts
    )

    prompt = f"""
You are a helpful assistant that answers
questions about the user's files.

RULES:

1. Use ONLY the supplied CONTEXT.
2. Do not use outside knowledge.
3. Do not guess.
4. Do not invent information.
5. If the answer is not in the context,
   say exactly:

"I couldn't find that information in the files."

6. Keep the answer concise.
7. Mention the source filename when possible.

CONTEXT:

{context}

QUESTION:

{question}

ANSWER:
"""

    start_time = time.perf_counter()

    response = HTTP_CLIENT.post(

        f"{OLLAMA_URL}/api/chat",

        json={

            "model":
                LLM_MODEL,

            "messages": [

                {
                    "role":
                        "system",

                    "content":
                        "Answer only from the supplied "
                        "document context. "
                        "Be concise."
                },

                {
                    "role":
                        "user",

                    "content":
                        prompt
                }
            ],

            # Return one complete response.
            "stream":
                False,

            # Disable Qwen3 thinking/reasoning
            # for faster document Q&A.
            "think":
                False,

            # Keep Qwen loaded for 10 minutes.
            # This is important for fast second questions.
            "keep_alive":
                LLM_KEEP_ALIVE,

            "options": {

                "temperature":
                    0,

                "num_ctx":
                    NUM_CTX,

                "num_predict":
                    NUM_PREDICT,

                "num_thread":
                    OLLAMA_NUM_THREADS
            }
        }
    )

    response.raise_for_status()

    data = response.json()

    # Debug information from Ollama
    prompt_eval_duration = data.get(
        "prompt_eval_duration",
        0
    )

    eval_duration = data.get(
        "eval_duration",
        0
    )

    prompt_eval_count = data.get(
        "prompt_eval_count",
        0
    )

    eval_count = data.get(
        "eval_count",
        0
    )

    print()
    print(
        f"Prompt tokens: {prompt_eval_count}"
    )

    print(
        f"Generated tokens: {eval_count}"
    )

    if prompt_eval_duration:
        print(
            f"Prompt processing: "
            f"{prompt_eval_duration / 1_000_000_000:.2f} seconds"
        )

    if eval_duration:
        print(
            f"Token generation: "
            f"{eval_duration / 1_000_000_000:.2f} seconds"
        )

    answer = (
        data
        .get("message", {})
        .get("content", "")
        .strip()
    )

    elapsed = (
            time.perf_counter()
            - start_time
    )

    print(
        f"AI generation completed in "
        f"{elapsed:.1f} seconds."
    )

    return answer


# ============================================================
# DISPLAY SEARCH RESULTS
# ============================================================

def display_search_results(results):

    print()

    print(
        "Top semantic search results:"
    )

    print()

    for number, result in enumerate(
            results,
            start=1
    ):

        print(

            f"{number}. "
            f"{result['filename']} "
            f"(similarity: "
            f"{result['score']:.3f})"

        )


# ============================================================
# BUILD OR LOAD INDEX
# ============================================================

def prepare_index(documents):

    documents_hash = calculate_documents_hash(
        documents
    )

    # --------------------------------------------------------
    # Try existing cache
    # --------------------------------------------------------

    cached_chunks, cached_embeddings = load_index(
        documents_hash
    )

    if (
            cached_chunks is not None
            and cached_embeddings is not None
    ):

        print()

        print(
            "Existing document index found."
        )

        print(
            f"Loaded {len(cached_chunks)} "
            f"cached chunks."
        )

        print(
            "Skipping embedding generation."
        )

        return (
            cached_chunks,
            cached_embeddings
        )

    # --------------------------------------------------------
    # Create chunks
    # --------------------------------------------------------

    print()

    print(
        "Splitting documents into chunks..."
    )

    chunks = create_chunks(
        documents
    )

    print(
        f"Created {len(chunks)} chunks."
    )

    if not chunks:

        raise RuntimeError(
            "No text chunks were created."
        )

    # --------------------------------------------------------
    # Create embeddings
    # --------------------------------------------------------

    try:

        embeddings = embed_chunks(
            chunks
        )

    except Exception as error:

        raise RuntimeError(
            "Could not create embeddings: "
            f"{error}"
        )

    # --------------------------------------------------------
    # Save index
    # --------------------------------------------------------

    print()

    print(
        "Saving document index..."
    )

    save_index(
        chunks,
        embeddings,
        documents_hash
    )

    print(
        "Document index saved."
    )

    return (
        chunks,
        embeddings
    )


# ============================================================
# MAIN CHAT
# ============================================================

def main():

    print(
        "==================================="
    )

    print(
        "       LOCAL AI FILE ASSISTANT"
    )

    print(
        "==================================="
    )

    print()

    # --------------------------------------------------------
    # CHECK OLLAMA
    # --------------------------------------------------------

    print(
        "Checking Ollama..."
    )

    if not check_ollama():

        print()

        print(
            "ERROR: Ollama is not running."
        )

        print()

        print(
            "Start Ollama and run this program again."
        )

        print()

        print(
            "You can test Ollama with:"
        )

        print(
            f"    ollama run {LLM_MODEL}"
        )

        return

    print(
        "Ollama is running."
    )

    # --------------------------------------------------------
    # CHECK MODELS
    # --------------------------------------------------------

    models = get_ollama_models()

    print()

    print(
        "Checking local models..."
    )

    # --------------------------------------------------------
    # LLM MODEL
    # --------------------------------------------------------

    if LLM_MODEL not in models:

        print()

        print(
            f"ERROR: LLM model "
            f"'{LLM_MODEL}' is not installed."
        )

        print()

        print(
            "Run:"
        )

        print(
            f"    ollama pull {LLM_MODEL}"
        )

        return

    # --------------------------------------------------------
    # EMBEDDING MODEL
    # --------------------------------------------------------

    if EMBEDDING_MODEL not in models:

        print()

        print(
            f"ERROR: Embedding model "
            f"'{EMBEDDING_MODEL}' "
            f"is not installed."
        )

        print()

        print(
            "Run:"
        )

        print(
            f"    ollama pull {EMBEDDING_MODEL}"
        )

        return

    print(
        f"LLM: {LLM_MODEL}"
    )

    print(
        f"LLM keep-alive: "
        f"{LLM_KEEP_ALIVE}"
    )

    print(
        f"Embeddings: {EMBEDDING_MODEL}"
    )

    print(
        f"Embedding keep-alive: "
        f"{EMBEDDING_KEEP_ALIVE}"
    )

    print(
        f"Context: {NUM_CTX}"
    )

    print(
        f"Max answer tokens: "
        f"{NUM_PREDICT}"
    )

    print(
        f"CPU threads: "
        f"{OLLAMA_NUM_THREADS}"
    )

    print(
        f"Top K: {TOP_K}"
    )

    print(
        f"Embedding batch size: "
        f"{EMBED_BATCH_SIZE}"
    )

    # --------------------------------------------------------
    # LOAD DOCUMENTS
    # --------------------------------------------------------

    print()

    print(
        "Loading files..."
    )

    documents = load_documents()

    print(
        f"Loaded {len(documents)} files."
    )

    if not documents:

        print()

        print(
            "No supported documents found."
        )

        print(
            "Supported formats:"
        )

        print(
            "    TXT"
        )

        print(
            "    CSV"
        )

        print(
            "    PDF"
        )

        print(
            "    DOCX"
        )

        print()

        print(
            "Put your files inside:"
        )

        print(
            f"    {DATA_FOLDER.absolute()}"
        )

        return

    # --------------------------------------------------------
    # BUILD OR LOAD INDEX
    # --------------------------------------------------------

    try:

        chunks, embeddings = prepare_index(
            documents
        )

    except Exception as error:

        print()

        print(
            "ERROR while preparing document index:"
        )

        print(
            error
        )

        return

    # --------------------------------------------------------
    # READY
    # --------------------------------------------------------

    print()

    print(
        "==================================="
    )

    print(
        "       DOCUMENT INDEX READY"
    )

    print(
        "==================================="
    )

    print()

    print(
        f"Files: {len(documents)}"
    )

    print(
        f"Chunks: {len(chunks)}"
    )

    print()

    print(
        "Ask questions about your files."
    )

    print(
        "Type 'exit' to quit."
    )

    print()

    # --------------------------------------------------------
    # CHAT LOOP
    # --------------------------------------------------------

    while True:

        try:

            question = input(
                "Ask a question: "
            ).strip()

        except (
                KeyboardInterrupt,
                EOFError
        ):

            print()

            print(
                "Goodbye!"
            )

            break

        if not question:

            continue

        if question.lower() == "exit":

            print(
                "Goodbye!"
            )

            break

        # ----------------------------------------------------
        # Greeting SEARCH
        # ----------------------------------------------------
        if is_greeting(question):

            print()
            print(
                "AI:"
            )
            print(
                "Hello! How can I help you with your files?"
            )
            print()

            continue

        # ----------------------------------------------------
        # SEMANTIC SEARCH
        # ----------------------------------------------------

        print()

        print(
            "Searching documents semantically..."
        )

        try:

            results = semantic_search(

                question,

                chunks,

                embeddings,

                TOP_K

            )

        except httpx.ConnectError:

            print()

            print(
                "Could not connect to Ollama."
            )

            print(
                "Make sure Ollama is running."
            )

            print()

            continue

        except Exception as error:

            print()

            print(
                "ERROR during semantic search:"
            )

            print(
                error
            )

            print()

            continue

        if not results:

            print()

            print(
                "I couldn't find relevant "
                "information in your files."
            )

            print()

            continue

        # ----------------------------------------------------
        # DISPLAY SEARCH RESULTS
        # ----------------------------------------------------

        display_search_results(
            results
        )

        # ----------------------------------------------------
        # ASK LLM
        # ----------------------------------------------------

        print()

        print(
            "Asking local AI..."
        )

        try:

            answer = ask_ai(

                question,

                results

            )

        except httpx.ConnectError:

            print()

            print(
                "Could not connect to Ollama."
            )

            print(
                "Make sure Ollama is running."
            )

            print()

            continue

        except httpx.HTTPStatusError as error:

            print()

            print(
                "Ollama returned an HTTP error:"
            )

            print(
                error
            )

            print()

            continue

        except Exception as error:

            print()

            print(
                "ERROR while asking local AI:"
            )

            print(
                error
            )

            print()

            continue

        # ----------------------------------------------------
        # DISPLAY ANSWER
        # ----------------------------------------------------

        print()

        print(
            "AI:"
        )

        print(
            answer
        )

        print()

        print(
            "-----------------------------------"
        )

        print()


# ============================================================
# PROGRAM ENTRY POINT
# ============================================================

if __name__ == "__main__":

    try:

        main()

    finally:

        HTTP_CLIENT.close()