from pathlib import Path
import hashlib
import json
import re
import time

import httpx
import numpy as np
import pandas as pd
from pypdf import PdfReader
from docx import Document


# ============================================================
# CONFIGURATION
# ============================================================

DATA_FOLDER = Path("../data")
INDEX_FOLDER = Path("../index")

OLLAMA_URL = "http://localhost:11434"

# ============================================================
# MODELS
# ============================================================

LLM_MODEL = "qwen3:1.7b"
EMBEDDING_MODEL = "nomic-embed-text:latest"

# ============================================================
# RAG SETTINGS
# ============================================================

TOP_K = 2

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150

# ============================================================
# PERFORMANCE SETTINGS
# ============================================================

EMBED_BATCH_SIZE = 2

# Keep these consistent with the Streamlit app.
NUM_CTX = 1536
NUM_PREDICT = 120

# Your machine is CPU-only.
OLLAMA_NUM_THREADS = 4

# Keep models loaded for a short period so repeated questions
# don't constantly reload them.
OLLAMA_KEEP_ALIVE = "5m"

# ============================================================
# CACHE FILES
# ============================================================

CHUNKS_FILE = INDEX_FOLDER / "chunks.json"
EMBEDDINGS_FILE = INDEX_FOLDER / "embeddings.npy"
CACHE_INFO_FILE = INDEX_FOLDER / "cache_info.json"


# ============================================================
# OLLAMA
# ============================================================

def check_ollama():
    """Check whether Ollama is running."""

    try:
        response = httpx.get(
            f"{OLLAMA_URL}/api/tags",
            timeout=5,
        )

        response.raise_for_status()

        return True

    except Exception:
        return False


def get_ollama_models():
    """Return installed Ollama model names."""

    try:
        response = httpx.get(
            f"{OLLAMA_URL}/api/tags",
            timeout=10,
        )

        response.raise_for_status()

        data = response.json()

        return [
            model["name"]
            for model in data.get("models", [])
        ]

    except Exception:
        return []


# ============================================================
# FILE READERS
# ============================================================

def read_txt(file_path):
    """Read TXT file."""

    with open(
            file_path,
            "r",
            encoding="utf-8",
            errors="ignore",
    ) as file:

        return file.read()


def read_csv(file_path):
    """Read CSV file."""

    dataframe = pd.read_csv(file_path)

    return dataframe.to_string(
        index=False
    )


def read_pdf(file_path):
    """Extract text from PDF."""

    reader = PdfReader(file_path)

    text = []

    for page_number, page in enumerate(
            reader.pages,
            start=1,
    ):

        page_text = page.extract_text()

        if page_text:

            text.append(
                f"[Page {page_number}]\n"
                f"{page_text}"
            )

    return "\n\n".join(text)


def read_docx(file_path):
    """Read DOCX paragraphs and tables."""

    document = Document(file_path)

    text = []

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        paragraph_text = paragraph.text.strip()

        if paragraph_text:

            text.append(
                paragraph_text
            )

    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    for table_number, table in enumerate(
            document.tables,
            start=1,
    ):

        text.append(
            f"\n[TABLE {table_number}]"
        )

        for row in table.rows:

            cells = []

            for cell in row.cells:

                cells.append(
                    cell.text.strip()
                )

            text.append(
                " | ".join(cells)
            )

    return "\n".join(text)


# ============================================================
# CONVERT FILE TO TEXT
# ============================================================

def convert_file_to_text(file_path):
    """Convert supported document into plain text."""

    extension = file_path.suffix.lower()

    try:

        if extension == ".txt":
            return read_txt(file_path)

        if extension == ".csv":
            return read_csv(file_path)

        if extension == ".pdf":
            return read_pdf(file_path)

        if extension == ".docx":
            return read_docx(file_path)

        return None

    except Exception as error:

        print(
            f"Could not read "
            f"{file_path.name}: {error}"
        )

        return None


# ============================================================
# LOAD DOCUMENTS
# ============================================================

def load_documents():
    """Load all supported files from data folder."""

    documents = []

    DATA_FOLDER.mkdir(
        parents=True,
        exist_ok=True,
    )

    for file_path in sorted(
            DATA_FOLDER.iterdir()
    ):

        if not file_path.is_file():
            continue

        if file_path.suffix.lower() not in {
            ".txt",
            ".csv",
            ".pdf",
            ".docx",
        }:
            continue

        text = convert_file_to_text(
            file_path
        )

        if text is None:
            continue

        text = text.strip()

        if not text:
            continue

        documents.append(
            {
                "filename": file_path.name,
                "content": text,
            }
        )

    return documents


# ============================================================
# TEXT CLEANING
# ============================================================

def clean_text(text):
    """Normalize whitespace."""

    text = re.sub(
        r"[ \t]+",
        " ",
        text,
    )

    text = re.sub(
        r"\n{3,}",
        "\n\n",
        text,
    )

    return text.strip()


# ============================================================
# CHUNK DOCUMENTS
# ============================================================

def create_chunks(documents):
    """Split documents into overlapping chunks."""

    chunks = []

    for document in documents:

        filename = document["filename"]

        content = clean_text(
            document["content"]
        )

        start = 0

        content_length = len(content)

        while start < content_length:

            end = min(
                start + CHUNK_SIZE,
                content_length,
                )

            chunk_text = content[
                start:end
            ].strip()

            if chunk_text:

                chunks.append(
                    {
                        "filename": filename,
                        "content": chunk_text,
                    }
                )

            if end >= content_length:
                break

            start = end - CHUNK_OVERLAP

    return chunks


# ============================================================
# DOCUMENT HASH
# ============================================================

def calculate_documents_hash(documents):
    """
    Calculate hash of documents + RAG configuration.

    If any document changes or chunking/embedding configuration
    changes, the cached index becomes invalid.
    """

    hasher = hashlib.sha256()

    hasher.update(
        str(CHUNK_SIZE).encode()
    )

    hasher.update(
        str(CHUNK_OVERLAP).encode()
    )

    hasher.update(
        EMBEDDING_MODEL.encode()
    )

    for document in documents:

        hasher.update(
            document["filename"].encode(
                "utf-8",
                errors="ignore",
            )
        )

        hasher.update(
            document["content"].encode(
                "utf-8",
                errors="ignore",
            )
        )

    return hasher.hexdigest()


# ============================================================
# SAVE INDEX
# ============================================================

def save_index(
        chunks,
        embeddings,
        documents_hash,
):
    """Save chunks, embeddings and cache metadata."""

    INDEX_FOLDER.mkdir(
        parents=True,
        exist_ok=True,
    )

    with open(
            CHUNKS_FILE,
            "w",
            encoding="utf-8",
    ) as file:

        json.dump(
            chunks,
            file,
            ensure_ascii=False,
            indent=2,
        )

    np.save(
        EMBEDDINGS_FILE,
        embeddings,
    )

    cache_info = {
        "documents_hash": documents_hash,
        "embedding_model": EMBEDDING_MODEL,
        "chunk_size": CHUNK_SIZE,
        "chunk_overlap": CHUNK_OVERLAP,
        "number_of_chunks": len(chunks),
    }

    with open(
            CACHE_INFO_FILE,
            "w",
            encoding="utf-8",
    ) as file:

        json.dump(
            cache_info,
            file,
            indent=2,
        )


# ============================================================
# LOAD INDEX
# ============================================================

def load_index(documents_hash):
    """Load cached index if it is still valid."""

    if not (
            CHUNKS_FILE.exists()
            and EMBEDDINGS_FILE.exists()
            and CACHE_INFO_FILE.exists()
    ):
        return None, None

    try:

        with open(
                CACHE_INFO_FILE,
                "r",
                encoding="utf-8",
        ) as file:

            cache_info = json.load(file)

        if (
                cache_info.get("documents_hash")
                != documents_hash
        ):
            return None, None

        if (
                cache_info.get("embedding_model")
                != EMBEDDING_MODEL
        ):
            return None, None

        if (
                cache_info.get("chunk_size")
                != CHUNK_SIZE
        ):
            return None, None

        if (
                cache_info.get("chunk_overlap")
                != CHUNK_OVERLAP
        ):
            return None, None

        with open(
                CHUNKS_FILE,
                "r",
                encoding="utf-8",
        ) as file:

            chunks = json.load(file)

        embeddings = np.load(
            EMBEDDINGS_FILE
        )

        if len(chunks) != len(embeddings):
            return None, None

        return chunks, embeddings

    except Exception as error:

        print(
            f"Could not load cache: {error}"
        )

        return None, None


# ============================================================
# OLLAMA EMBEDDINGS
# ============================================================

def create_embeddings(texts):
    """Create embeddings using Ollama."""

    response = httpx.post(
        f"{OLLAMA_URL}/api/embed",
        json={
            "model": EMBEDDING_MODEL,
            "input": texts,
            "keep_alive": OLLAMA_KEEP_ALIVE,
        },
        timeout=300,
    )

    response.raise_for_status()

    data = response.json()

    embeddings = data.get(
        "embeddings"
    )

    if not embeddings:

        raise RuntimeError(
            "Ollama did not return embeddings."
        )

    return np.array(
        embeddings,
        dtype=np.float32,
    )


# ============================================================
# CREATE EMBEDDINGS
# ============================================================

def embed_chunks(
        chunks,
        progress_callback=None,
):
    """Create embeddings for document chunks."""

    texts = [
        chunk["content"]
        for chunk in chunks
    ]

    embeddings = []

    total = len(texts)

    for start in range(
            0,
            total,
            EMBED_BATCH_SIZE,
    ):

        batch = texts[
            start:start + EMBED_BATCH_SIZE
        ]

        end_number = min(
            start + EMBED_BATCH_SIZE,
            total,
            )

        if progress_callback:

            progress_callback(
                f"Creating embeddings "
                f"{end_number}/{total}..."
            )

        batch_embeddings = create_embeddings(
            batch
        )

        embeddings.extend(
            batch_embeddings
        )

    return np.array(
        embeddings,
        dtype=np.float32,
    )


# ============================================================
# COSINE SIMILARITY
# ============================================================

def cosine_similarity(a, b):
    """Calculate cosine similarity."""

    denominator = (
            np.linalg.norm(a)
            *
            np.linalg.norm(b)
    )

    if denominator == 0:
        return 0.0

    return float(
        np.dot(a, b)
        / denominator
    )


# ============================================================
# SEMANTIC SEARCH
# ============================================================

def semantic_search(
        question,
        chunks,
        embeddings,
        top_k=TOP_K,
        selected_files=None,
):
    """
    Search document chunks by semantic similarity.

    selected_files=None means search all documents.
    """

    question_embedding = create_embeddings(
        [question]
    )[0]

    scores = []

    for index, embedding in enumerate(
            embeddings
    ):

        filename = chunks[index]["filename"]

        if (
                selected_files
                and filename not in selected_files
        ):
            continue

        score = cosine_similarity(
            question_embedding,
            embedding,
        )

        scores.append(
            (
                score,
                index,
            )
        )

    scores.sort(
        key=lambda item: item[0],
        reverse=True,
    )

    results = []

    for score, index in scores[:top_k]:

        results.append(
            {
                "filename":
                    chunks[index]["filename"],

                "content":
                    chunks[index]["content"],

                "score":
                    score,
            }
        )

    return results


# ============================================================
# ASK LOCAL LLM
# ============================================================

def ask_ai(
        question,
        search_results,
):
    """
    Ask Qwen using ONLY retrieved context.

    Important:
    The LLM is NOT responsible for formatting sources.
    Python handles source display separately.
    """

    context_parts = []

    for number, result in enumerate(
            search_results,
            start=1,
    ):

        context_parts.append(
            f"DOCUMENT {number}\n"
            f"FILE: {result['filename']}\n"
            f"CONTENT:\n{result['content']}"
        )

    context = "\n\n".join(
        context_parts
    )

    prompt = f"""
You are a document question-answering assistant.

Answer the user's question using ONLY the supplied
document context.

RULES:

1. Use ONLY the CONTEXT.
2. Do not use outside knowledge.
3. Do not guess.
4. Do not invent information.
5. Give a direct answer.
6. Use complete natural sentences.
7. Do not output filenames.
8. Do not output source labels.
9. Do not output "SOURCE:".
10. Do not explain your reasoning.
11. If the answer is not contained in the context,
    reply exactly:

I couldn't find that information in the files.

CONTEXT:

{context}

QUESTION:

{question}

ANSWER:
""".strip()

    start_time = time.perf_counter()

    response = httpx.post(
        f"{OLLAMA_URL}/api/chat",
        json={
            "model": LLM_MODEL,

            "messages": [
                {
                    "role": "system",
                    "content": (
                        "You answer questions using only "
                        "the supplied document context. "
                        "Return only the concise answer."
                    ),
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],

            "stream": False,

            # Disable Qwen thinking for faster/simple answers.
            "think": False,

            # Keep Qwen loaded for subsequent questions.
            "keep_alive": OLLAMA_KEEP_ALIVE,

            "options": {
                "temperature": 0,
                "num_ctx": NUM_CTX,
                "num_predict": NUM_PREDICT,
                "num_thread": OLLAMA_NUM_THREADS,
            },
        },
        timeout=300,
    )

    response.raise_for_status()

    total_time = (
            time.perf_counter()
            - start_time
    )

    data = response.json()

    answer = (
        data
        .get("message", {})
        .get("content", "")
        .strip()
    )

    # --------------------------------------------------------
    # Remove accidental source formatting from small models.
    # --------------------------------------------------------

    answer = re.sub(
        r"\s*/\s*[\w.\- ]+\.(csv|txt|pdf|docx)\s*$",
        "",
        answer,
        flags=re.IGNORECASE,
    ).strip()

    prompt_tokens = data.get(
        "prompt_eval_count",
        0,
    )

    generated_tokens = data.get(
        "eval_count",
        0,
    )

    prompt_duration = data.get(
        "prompt_eval_duration",
        0,
    )

    generation_duration = data.get(
        "eval_duration",
        0,
    )

    return {
        "answer": answer,

        "prompt_tokens":
            prompt_tokens,

        "generated_tokens":
            generated_tokens,

        "prompt_time":
            prompt_duration / 1_000_000_000,

        "generation_time":
            generation_duration / 1_000_000_000,

        "total_time":
            total_time,
    }


# ============================================================
# BUILD OR LOAD INDEX
# ============================================================

def prepare_index(
        documents,
        progress_callback=None,
        force_rebuild=False,
):
    """
    Load cached index or rebuild it.
    """

    documents_hash = calculate_documents_hash(
        documents
    )

    # --------------------------------------------------------
    # Try cache first
    # --------------------------------------------------------

    if not force_rebuild:

        cached_chunks, cached_embeddings = (
            load_index(documents_hash)
        )

        if (
                cached_chunks is not None
                and cached_embeddings is not None
        ):

            if progress_callback:

                progress_callback(
                    f"Loaded cached index: "
                    f"{len(cached_chunks)} chunks"
                )

            return (
                cached_chunks,
                cached_embeddings,
                True,
            )

    # --------------------------------------------------------
    # Create chunks
    # --------------------------------------------------------

    if progress_callback:

        progress_callback(
            "Creating document chunks..."
        )

    chunks = create_chunks(
        documents
    )

    if not chunks:

        raise RuntimeError(
            "No text chunks were created."
        )

    # --------------------------------------------------------
    # Create embeddings
    # --------------------------------------------------------

    embeddings = embed_chunks(
        chunks,
        progress_callback,
    )

    # --------------------------------------------------------
    # Save
    # --------------------------------------------------------

    if progress_callback:

        progress_callback(
            "Saving document index..."
        )

    save_index(
        chunks,
        embeddings,
        documents_hash,
    )

    if progress_callback:

        progress_callback(
            "Document index saved."
        )

    return (
        chunks,
        embeddings,
        False,
    )


# ============================================================
# INITIALIZE ENGINE
# ============================================================

def initialize(
        progress_callback=None,
        force_rebuild=False,
):
    """
    Initialize the complete RAG engine.
    """

    # --------------------------------------------------------
    # Ollama
    # --------------------------------------------------------

    if not check_ollama():

        raise RuntimeError(
            "Ollama is not running."
        )

    # --------------------------------------------------------
    # Models
    # --------------------------------------------------------

    models = get_ollama_models()

    if LLM_MODEL not in models:

        raise RuntimeError(
            f"LLM model '{LLM_MODEL}' "
            f"is not installed.\n\n"
            f"Run:\n"
            f"ollama pull {LLM_MODEL}"
        )

    if EMBEDDING_MODEL not in models:

        raise RuntimeError(
            f"Embedding model '{EMBEDDING_MODEL}' "
            f"is not installed.\n\n"
            f"Run:\n"
            f"ollama pull {EMBEDDING_MODEL}"
        )

    # --------------------------------------------------------
    # Documents
    # --------------------------------------------------------

    if progress_callback:

        progress_callback(
            "Loading documents..."
        )

    documents = load_documents()

    if not documents:

        raise RuntimeError(
            "No supported documents found "
            "in the data folder."
        )

    if progress_callback:

        progress_callback(
            f"Loaded {len(documents)} files."
        )

    # --------------------------------------------------------
    # Index
    # --------------------------------------------------------

    (
        chunks,
        embeddings,
        cache_used,
    ) = prepare_index(
        documents,
        progress_callback,
        force_rebuild=force_rebuild,
    )

    return {
        "documents":
            documents,

        "chunks":
            chunks,

        "embeddings":
            embeddings,

        "cache_used":
            cache_used,
    }


# ============================================================
# ASK QUESTION
# ============================================================

def ask_question(
        question,
        engine,
        selected_files=None,
):
    """
    Complete RAG pipeline:

    Question
       ↓
    Embedding
       ↓
    Semantic search
       ↓
    Retrieved context
       ↓
    Qwen
       ↓
    Answer
    """

    search_start = time.perf_counter()

    results = semantic_search(
        question,
        engine["chunks"],
        engine["embeddings"],
        TOP_K,
        selected_files,
    )

    search_time = (
            time.perf_counter()
            - search_start
    )

    if not results:

        return {
            "answer":
                "I couldn't find that information "
                "in the files.",

            "sources":
                [],

            "search_time":
                search_time,

            "total_time":
                search_time,
        }

    llm_result = ask_ai(
        question,
        results,
    )

    return {
        "answer":
            llm_result["answer"],

        "sources":
            results,

        "search_time":
            search_time,

        "prompt_tokens":
            llm_result["prompt_tokens"],

        "generated_tokens":
            llm_result["generated_tokens"],

        "prompt_time":
            llm_result["prompt_time"],

        "generation_time":
            llm_result["generation_time"],

        "ai_time":
            llm_result["total_time"],

        "total_time":
            search_time
            +
            llm_result["total_time"],
    }