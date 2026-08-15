from pathlib import Path
import hashlib
import json
import re
import time

import httpx
import numpy as np
import pandas as pd
from pypdf import PdfReader
from docx import Document


# ============================================================
# CONFIGURATION
# ============================================================

DATA_FOLDER = Path("../data")
INDEX_FOLDER = Path("../index")

OLLAMA_URL = "http://localhost:11434"

# Local LLM
LLM_MODEL = "qwen3:1.7b"

# Local embedding model
EMBEDDING_MODEL = "nomic-embed-text:latest"

# ============================================================
# RAG SETTINGS
# ============================================================

TOP_K = 2

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150

# ============================================================
# PERFORMANCE SETTINGS
# ============================================================

EMBED_BATCH_SIZE = 2

NUM_CTX = 1536
NUM_PREDICT = 250

OLLAMA_NUM_THREADS = 2

# Keep embedding model alive briefly.
EMBED_KEEP_ALIVE = "5m"

# We deliberately unload the LLM after every answer
# because your machine has limited RAM.
LLM_KEEP_ALIVE = 0

# ============================================================
# CACHE FILES
# ============================================================

CHUNKS_FILE = INDEX_FOLDER / "chunks.json"
EMBEDDINGS_FILE = INDEX_FOLDER / "embeddings.npy"
CACHE_INFO_FILE = INDEX_FOLDER / "cache_info.json"


# ============================================================
# OLLAMA
# ============================================================

def check_ollama():
    try:
        response = httpx.get(
            f"{OLLAMA_URL}/api/tags",
            timeout=5
        )

        response.raise_for_status()

        return True

    except Exception:
        return False


def get_ollama_models():
    try:
        response = httpx.get(
            f"{OLLAMA_URL}/api/tags",
            timeout=10
        )

        response.raise_for_status()

        data = response.json()

        return [
            model["name"]
            for model in data.get("models", [])
        ]

    except Exception:
        return []


# ============================================================
# FILE READERS
# ============================================================

def read_txt(file_path):
    with open(
            file_path,
            "r",
            encoding="utf-8",
            errors="ignore"
    ) as file:
        return file.read()


def read_csv(file_path):
    dataframe = pd.read_csv(file_path)

    return dataframe.to_string(
        index=False
    )


def read_pdf(file_path):
    reader = PdfReader(file_path)

    text = []

    for page_number, page in enumerate(
            reader.pages,
            start=1
    ):
        page_text = page.extract_text()

        if page_text:
            text.append(
                f"[Page {page_number}]\n"
                f"{page_text}"
            )

    return "\n\n".join(text)


def read_docx(file_path):
    document = Document(file_path)

    text = []

    # Paragraphs
    for paragraph in document.paragraphs:

        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            text.append(paragraph_text)

    # Tables
    for table_number, table in enumerate(
            document.tables,
            start=1
    ):

        text.append(
            f"\n[TABLE {table_number}]"
        )

        for row in table.rows:

            cells = []

            for cell in row.cells:
                cells.append(
                    cell.text.strip()
                )

            text.append(
                " | ".join(cells)
            )

    return "\n".join(text)


# ============================================================
# CONVERT FILE TO TEXT
# ============================================================

def convert_file_to_text(file_path):

    extension = file_path.suffix.lower()

    try:

        if extension == ".txt":
            return read_txt(file_path)

        elif extension == ".csv":
            return read_csv(file_path)

        elif extension == ".pdf":
            return read_pdf(file_path)

        elif extension == ".docx":
            return read_docx(file_path)

        return None

    except Exception as error:

        print(
            f"Could not read "
            f"{file_path.name}: {error}"
        )

        return None


# ============================================================
# LOAD DOCUMENTS
# ============================================================

def load_documents():

    documents = []

    if not DATA_FOLDER.exists():

        print(
            "ERROR: Data folder does not exist:"
        )

        print(
            f"    {DATA_FOLDER.absolute()}"
        )

        return documents

    for file_path in sorted(
            DATA_FOLDER.iterdir()
    ):

        if not file_path.is_file():
            continue

        text = convert_file_to_text(
            file_path
        )

        if text is None:
            continue

        text = text.strip()

        if not text:
            continue

        documents.append({
            "filename": file_path.name,
            "content": text
        })

    return documents


# ============================================================
# TEXT CLEANING
# ============================================================

def clean_text(text):

    text = re.sub(
        r"[ \t]+",
        " ",
        text
    )

    text = re.sub(
        r"\n{3,}",
        "\n\n",
        text
    )

    return text.strip()


# ============================================================
# CHUNK DOCUMENTS
# ============================================================

def create_chunks(documents):

    chunks = []

    for document in documents:

        filename = document["filename"]

        content = clean_text(
            document["content"]
        )

        start = 0

        content_length = len(content)

        while start < content_length:

            end = min(
                start + CHUNK_SIZE,
                content_length
            )

            chunk_text = content[
                start:end
            ].strip()

            if chunk_text:

                chunks.append({
                    "filename": filename,
                    "content": chunk_text
                })

            if end >= content_length:
                break

            start = end - CHUNK_OVERLAP

    return chunks


# ============================================================
# CACHE / INDEX
# ============================================================

def calculate_documents_hash(documents):

    hasher = hashlib.sha256()

    hasher.update(
        str(CHUNK_SIZE).encode()
    )

    hasher.update(
        str(CHUNK_OVERLAP).encode()
    )

    hasher.update(
        EMBEDDING_MODEL.encode()
    )

    for document in documents:

        hasher.update(
            document["filename"].encode(
                "utf-8",
                errors="ignore"
            )
        )

        hasher.update(
            document["content"].encode(
                "utf-8",
                errors="ignore"
            )
        )

    return hasher.hexdigest()


def save_index(
        chunks,
        embeddings,
        documents_hash
):

    INDEX_FOLDER.mkdir(
        parents=True,
        exist_ok=True
    )

    with open(
            CHUNKS_FILE,
            "w",
            encoding="utf-8"
    ) as file:

        json.dump(
            chunks,
            file,
            ensure_ascii=False,
            indent=2
        )

    np.save(
        EMBEDDINGS_FILE,
        embeddings
    )

    cache_info = {
        "documents_hash": documents_hash,
        "embedding_model": EMBEDDING_MODEL,
        "chunk_size": CHUNK_SIZE,
        "chunk_overlap": CHUNK_OVERLAP,
        "number_of_chunks": len(chunks)
    }

    with open(
            CACHE_INFO_FILE,
            "w",
            encoding="utf-8"
    ) as file:

        json.dump(
            cache_info,
            file,
            indent=2
        )


def load_index(documents_hash):

    if not (
            CHUNKS_FILE.exists()
            and EMBEDDINGS_FILE.exists()
            and CACHE_INFO_FILE.exists()
    ):
        return None, None

    try:

        with open(
                CACHE_INFO_FILE,
                "r",
                encoding="utf-8"
        ) as file:

            cache_info = json.load(file)

        if (
                cache_info.get("documents_hash")
                != documents_hash
        ):
            return None, None

        if (
                cache_info.get("embedding_model")
                != EMBEDDING_MODEL
        ):
            return None, None

        if (
                cache_info.get("chunk_size")
                != CHUNK_SIZE
        ):
            return None, None

        if (
                cache_info.get("chunk_overlap")
                != CHUNK_OVERLAP
        ):
            return None, None

        with open(
                CHUNKS_FILE,
                "r",
                encoding="utf-8"
        ) as file:

            chunks = json.load(file)

        embeddings = np.load(
            EMBEDDINGS_FILE
        )

        if len(chunks) != len(embeddings):
            return None, None

        return chunks, embeddings

    except Exception:

        return None, None


# ============================================================
# OLLAMA EMBEDDINGS
# ============================================================

def create_embeddings(texts):

    response = httpx.post(

        f"{OLLAMA_URL}/api/embed",

        json={
            "model": EMBEDDING_MODEL,
            "input": texts,
            "keep_alive": EMBED_KEEP_ALIVE
        },

        timeout=300
    )

    response.raise_for_status()

    data = response.json()

    embeddings = data.get(
        "embeddings"
    )

    if not embeddings:

        raise RuntimeError(
            "Ollama did not return embeddings."
        )

    return np.array(
        embeddings,
        dtype=np.float32
    )


def embed_chunks(chunks):

    texts = [
        chunk["content"]
        for chunk in chunks
    ]

    embeddings = []

    for start in range(
            0,
            len(texts),
            EMBED_BATCH_SIZE
    ):

        batch = texts[
            start:
            start + EMBED_BATCH_SIZE
        ]

        batch_embeddings = create_embeddings(
            batch
        )

        embeddings.extend(
            batch_embeddings
        )

    return np.array(
        embeddings,
        dtype=np.float32
    )


# ============================================================
# BUILD OR LOAD INDEX
# ============================================================

def prepare_index(documents):

    documents_hash = calculate_documents_hash(
        documents
    )

    cached_chunks, cached_embeddings = load_index(
        documents_hash
    )

    if (
            cached_chunks is not None
            and cached_embeddings is not None
    ):

        return (
            cached_chunks,
            cached_embeddings,
            True
        )

    chunks = create_chunks(
        documents
    )

    if not chunks:

        raise RuntimeError(
            "No text chunks were created."
        )

    embeddings = embed_chunks(
        chunks
    )

    save_index(
        chunks,
        embeddings,
        documents_hash
    )

    return (
        chunks,
        embeddings,
        False
    )


# ============================================================
# COSINE SIMILARITY
# ============================================================

def semantic_search(
        question,
        chunks,
        embeddings,
        top_k=TOP_K
):

    start_time = time.perf_counter()

    question_embedding = create_embeddings(
        [question]
    )[0]

    # Vectorized cosine similarity.
    # Faster than calculating each chunk separately.
    embedding_norms = np.linalg.norm(
        embeddings,
        axis=1
    )

    question_norm = np.linalg.norm(
        question_embedding
    )

    denominators = (
            embedding_norms
            * question_norm
    )

    similarities = np.zeros(
        len(embeddings),
        dtype=np.float32
    )

    valid = denominators != 0

    similarities[valid] = (
            np.dot(
                embeddings[valid],
                question_embedding
            )
            / denominators[valid]
    )

    top_indices = np.argsort(
        similarities
    )[::-1][:top_k]

    results = []

    for index in top_indices:

        results.append({

            "filename":
                chunks[index]["filename"],

            "content":
                chunks[index]["content"],

            "score":
                float(similarities[index])
        })

    elapsed = time.perf_counter() - start_time

    return results, elapsed


# ============================================================
# ASK LOCAL LLM
# ============================================================

def ask_ai(
        question,
        search_results
):

    context_parts = []

    for result in search_results:

        context_parts.append(

            "SOURCE: "
            + result["filename"]
            + "\n"
            + result["content"]

        )

    context = "\n\n".join(
        context_parts
    )

    prompt = f"""
You are a document question-answering assistant.

Use ONLY the supplied CONTEXT.

Rules:
1. Do not use outside knowledge.
2. Do not guess.
3. Do not invent information.
4. Answer directly and concisely.
5. If the answer is not present in the context,
   say exactly:
"I couldn't find that information in the files."
6. Include the source filename when useful.

CONTEXT:

{context}

QUESTION:

{question}

ANSWER:
"""

    start_time = time.perf_counter()

    response = httpx.post(

        f"{OLLAMA_URL}/api/chat",

        json={

            "model": LLM_MODEL,

            "messages": [

                {
                    "role": "system",
                    "content":
                        "Answer only from the supplied "
                        "document context. "
                        "Be concise."
                },

                {
                    "role": "user",
                    "content": prompt
                }

            ],

            "stream": False,

            "think": False,

            "keep_alive": LLM_KEEP_ALIVE,

            "options": {

                "temperature": 0,

                "num_ctx": NUM_CTX,

                "num_predict": NUM_PREDICT,

                "num_thread": OLLAMA_NUM_THREADS
            }
        },

        timeout=300
    )

    response.raise_for_status()

    data = response.json()

    answer = (
        data
        .get("message", {})
        .get("content", "")
        .strip()
    )

    elapsed = time.perf_counter() - start_time

    # Ollama performance information
    prompt_tokens = data.get(
        "prompt_eval_count"
    )

    generated_tokens = data.get(
        "eval_count"
    )

    return {
        "answer": answer,
        "elapsed": elapsed,
        "prompt_tokens": prompt_tokens,
        "generated_tokens": generated_tokens
    }


# ============================================================
# INITIALIZE BACKEND
# ============================================================

def initialize():

    if not check_ollama():

        raise RuntimeError(
            "Ollama is not running. "
            "Start Ollama first."
        )

    models = get_ollama_models()

    if LLM_MODEL not in models:

        raise RuntimeError(
            f"LLM model '{LLM_MODEL}' "
            "is not installed."
        )

    if EMBEDDING_MODEL not in models:

        raise RuntimeError(
            f"Embedding model '{EMBEDDING_MODEL}' "
            "is not installed."
        )

    documents = load_documents()

    if not documents:

        raise RuntimeError(
            "No supported documents found "
            "inside the data folder."
        )

    chunks, embeddings, cache_used = prepare_index(
        documents
    )

    return {
        "documents": documents,
        "chunks": chunks,
        "embeddings": embeddings,
        "cache_used": cache_used
    }